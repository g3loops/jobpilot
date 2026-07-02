"""Tailored résumé + cover-letter generation (signature deliverable + cover-letter bonus).

Pipeline:
  1. Claude (Haiku) tailors the candidate's OWN resume text to the selected job, returning
     STRUCTURED content (summary, experience entries, projects, grouped skills, education).
     It only re-orders / re-phrases existing experience — never invents employers or dates.
     Falls back to a deterministic template build if no API key / the call fails.
  2. The same structured data is rendered into BOTH:
       * a Word .docx  (python-docx)   — editable
       * a PDF         (fpdf2)         — pure-Python, Cloud-Run friendly
     in a clean single-column style (centred navy name, ruled section headers, "Title |
     Company" left with dates right-aligned, grouped skills) that matches the user's template.
The cover letter mirrors this in a serif (Times) style.
"""
from __future__ import annotations

import io
import json
import re

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from . import config

NAVY = (0x1F, 0x4E, 0x79)
GRAY = (0x55, 0x55, 0x55)


# ===========================================================================
# Helpers
# ===========================================================================
def _matched_missing(job: pd.Series, prof):
    job_sk = set(job.get("skills", []) or [])
    prof_sk = set(prof.skills)
    return sorted(job_sk & prof_sk), sorted(prof_sk - job_sk)


def _pretty_skill(s: str) -> str:
    return s.title() if len(s) > 3 else s.upper()


def _latin1(text) -> str:
    """fpdf2 core fonts are latin-1 only; map the unicode we emit to safe equivalents."""
    if text is None:
        return ""
    repl = {"—": "-", "–": "-", "•": "-", "✓": "", "→": "->",
            "’": "'", "‘": "'", "“": '"', "”": '"', "…": "...", " ": " "}
    for a, b in repl.items():
        text = str(text).replace(a, b)
    return text.encode("latin-1", "replace").decode("latin-1")


# ===========================================================================
# Structured content (Claude-tailored, with deterministic fallback)
# ===========================================================================
RESUME_SCHEMA_HINT = (
    '{\n'
    '  "summary": "2-3 sentence professional summary tailored to the target job",\n'
    '  "experience": [{"title": "", "company": "", "location": "", "dates": "", "bullets": ["", ""]}],\n'
    '  "projects":   [{"title": "", "company": "", "location": "", "dates": "", "bullets": [""]}],\n'
    '  "publications": ["citation string", "..."],\n'
    '  "skills":     {"Tools": [""], "Finance": [""], "Analytics": [""]},\n'
    '  "education":  [{"degree": "", "school": "", "dates": "", "coursework": ""}]\n'
    '}'
)


def _llm_resume_json(job: pd.Series, prof) -> dict | None:
    """Ask Claude (Haiku) to tailor the candidate's resume to this job, as structured JSON.
    Returns None on any failure so the caller falls back to the deterministic build."""
    if not config.ANTHROPIC_API_KEY:
        return None
    try:
        import anthropic
    except ImportError:
        return None
    try:
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        jd = str(job.get("description", ""))[:2500]
        title, company = job.get("title", "the role"), job.get("company", "the company")
        system = (
            "You are an expert résumé writer. Return ONLY valid JSON matching the requested "
            "schema — no markdown fences, no prose. TAILOR the candidate's EXISTING resume to "
            "the target job: reorder and rephrase bullets to surface the most relevant "
            "experience, and write a fitted summary. NEVER invent employers, job titles, dates, "
            "degrees, or skills the candidate does not already have. Omit a section if the "
            "candidate has nothing for it (empty list). Keep bullets concise and impact-first."
        )
        hints = []
        emphasis = [e for e in (getattr(prof, "emphasis", []) or []) if e]
        if emphasis:
            hints.append(
                f"LEAD the summary and the Skills section with these (in this order): "
                f"{', '.join(emphasis)}. The FIRST sentence of the summary must foreground "
                f"them, and the Skills section must list them first. Reframe the candidate's "
                f"prior experience to demonstrate these as core strengths — e.g. present "
                f"Kafka / Spark / Kubernetes as ML-infrastructure & MLOps platform work "
                f"(feature pipelines, distributed training, model serving), not generic "
                f"backend; and surface Python + machine learning ahead of reporting/analytics "
                f"for an ML pivot. Never claim skills the candidate lacks.")
        if getattr(prof, "education_first", False):
            hints.append("The candidate is a new graduate — keep work experience concise; the "
                         "résumé will LEAD with Education, so make the summary education-forward.")
        if getattr(prof, "publications", []):
            hints.append("Populate the 'publications' array with the candidate's publications "
                         "(verbatim) — they should be featured prominently.")
        if prof.target_roles:
            hints.append(f"Frame transferable skills toward the target role(s) "
                         f"({', '.join(prof.target_roles)}): lead with the skills the job "
                         f"screens for rather than generic reporting.")
        pubs = "\n".join(getattr(prof, "publications", []) or [])
        ask = (
            f"Tailor this candidate's résumé to the target job. Return JSON exactly in this "
            f"schema:\n{RESUME_SCHEMA_HINT}\n\n"
            f"CANDIDATE NAME: {prof.name}\n"
            f"KNOWN SKILLS: {', '.join(prof.skills)}\n"
            f"TARGET ROLES: {', '.join(prof.target_roles)}\n"
            + (f"KNOWN PUBLICATIONS:\n{pubs}\n" if pubs else "")
            + (f"\nGUIDANCE: {' '.join(hints)}\n" if hints else "")
            + f"\nCANDIDATE RÉSUMÉ (source of truth — do not go beyond it):\n{prof.resume_text}\n\n"
            f"TARGET JOB: {title} at {company}\n"
            f"JOB DESCRIPTION:\n{jd}"
        )
        # Haiku only for this app — never Opus/Sonnet.
        msg = client.messages.create(
            model="claude-haiku-4-5", max_tokens=2500, system=system,
            messages=[{"role": "user", "content": ask}],
        )
        raw = "".join(b.text for b in msg.content if b.type == "text").strip()
        raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
        raw = raw[raw.find("{"): raw.rfind("}") + 1]   # isolate the JSON object
        data = json.loads(raw)
        if not isinstance(data, dict) or "summary" not in data:
            return None
        return data
    except Exception:
        return None


def _entry(e) -> dict:
    """Normalise an experience/project entry (dict or free-text string) to a common shape."""
    if isinstance(e, dict):
        return {"title": e.get("title", ""), "company": e.get("company", e.get("org", "")),
                "location": e.get("location", ""), "dates": e.get("dates", ""),
                "bullets": [b for b in e.get("bullets", []) if b]}
    s = str(e)
    # "Title — description"  ->  title + one bullet
    m = re.split(r"\s[—–-]\s", s, maxsplit=1)
    if len(m) == 2:
        return {"title": m[0].strip(), "company": "", "location": "", "dates": "",
                "bullets": [m[1].strip()]}
    return {"title": "", "company": "", "location": "", "dates": "", "bullets": [s]}


def _edu_entry(e) -> dict:
    if isinstance(e, dict):
        return {"degree": e.get("degree", ""), "school": e.get("school", ""),
                "dates": e.get("dates", ""), "coursework": e.get("coursework", "")}
    return {"degree": str(e), "school": "", "dates": "", "coursework": ""}


def build_resume_data(job: pd.Series, prof, use_llm: bool = True) -> dict:
    """Produce the structured résumé (Claude-tailored if possible, else template) + ATS info."""
    from .ats import ats_match
    ats = ats_match(job, prof)
    matched, _ = _matched_missing(job, prof)
    title, company = str(job.get("title", "the role")), str(job.get("company", "the company"))

    llm = _llm_resume_json(job, prof) if use_llm else None
    backend = "llm" if llm else "template"

    if llm:
        experience = [_entry(e) for e in llm.get("experience", []) if e]
        projects = [_entry(e) for e in llm.get("projects", []) if e]
        education = [_edu_entry(e) for e in llm.get("education", []) if e]
        skills = llm.get("skills") or {}
        if isinstance(skills, list):
            skills = {"Skills": skills}
        summary = llm.get("summary", "")
    else:
        experience = [_entry(e) for e in (prof.experience or [])]
        projects = [_entry(e) for e in (getattr(prof, "projects", []) or [])]
        education = [_edu_entry(e) for e in (prof.education or [])]
        # Order skills so the emphasis themes lead, then matched, then the rest.
        emphasis = [e.lower() for e in (getattr(prof, "emphasis", []) or [])]

        def _skill_rank(s: str):
            sl = s.lower()
            for i, e in enumerate(emphasis):
                if e in sl or sl in e:
                    return (0, i)
            return (1, 0 if s in matched else 1)

        lead = sorted(prof.skills, key=_skill_rank)
        skills = {"Skills": [_pretty_skill(s) for s in lead]}
        roles = ", ".join(prof.target_roles[:3]) if prof.target_roles else title
        emph_lead = ", ".join(getattr(prof, "emphasis", [])[:3])
        top = (ats["covered"] or matched or prof.skills)[:5]
        summary = (
            (f"{emph_lead}-focused " if emph_lead else "")
            + f"{prof.target_roles[0] if prof.target_roles else 'Analytics professional'} "
            f"targeting {roles}. Hands-on strength in {emph_lead or ', '.join(top)}. Strong "
            f"fit for the {title} role at {company}, bringing directly relevant experience "
            f"and a track record of turning data into decisions."
        )

    # clean empty skill groups
    skills = {k: [s for s in v if s] for k, v in skills.items() if v and any(v)}

    publications = ([p for p in (llm.get("publications") or []) if p] if llm else [])
    if not publications:
        publications = list(getattr(prof, "publications", []) or [])
    education_first = bool(getattr(prof, "education_first", False))
    skills_first = bool(getattr(prof, "skills_first", False))
    # Section order: publications lead (if any); new grads lead with Education; career
    # pivoters (skills_first) lead with Skills so Python/ML sits right under the summary.
    order = ["profile"] + (["publications"] if publications else [])
    if education_first:
        order += ["education", "experience", "projects", "skills"]
    elif skills_first:
        order += ["skills", "experience", "projects", "education"]
    else:
        order += ["experience", "projects", "skills", "education"]

    return {"name": prof.name, "contact": prof.contact, "summary": summary,
            "experience": experience, "projects": projects, "skills": skills,
            "education": education, "publications": publications,
            "section_order": order, "ats": ats, "backend": backend,
            "title": title, "company": company}


def build_cover_data(job: pd.Series, prof, use_llm: bool = True) -> dict:
    title, company = str(job.get("title", "the role")), str(job.get("company", "your team"))
    paras = _llm_cover_paragraphs(job, prof) if use_llm else None
    backend = "llm" if paras else "template"
    if not paras:
        matched, _ = _matched_missing(job, prof)
        emphasis = [e for e in (getattr(prof, "emphasis", []) or []) if e]
        skills_phrase = (", ".join(emphasis[:5]) if emphasis
                         else ", ".join(matched[:5]) if matched else ", ".join(prof.skills[:5]))
        paras = [
            f"I am excited to apply for the {title} position at {company}. As "
            f"{prof.target_roles[0] if prof.target_roles else 'a candidate'} with hands-on "
            f"experience in {skills_phrase}, I am confident I can contribute from day one.",
            prof.resume_text.strip(),
            f"What draws me to this role is the alignment between your requirements and my "
            f"strengths in {', '.join(matched[:4]) if matched else 'the core skills you list'}. "
            f"I would welcome the opportunity to discuss how I can help {company} succeed.",
        ]
    return {"name": prof.name, "contact": prof.contact, "greeting": "Dear Hiring Manager,",
            "paragraphs": [p for p in paras if p and p.strip()],
            "closing": "Sincerely,", "backend": backend}


def _llm_cover_paragraphs(job: pd.Series, prof) -> list | None:
    if not config.ANTHROPIC_API_KEY:
        return None
    try:
        import anthropic
    except ImportError:
        return None
    try:
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        jd = str(job.get("description", ""))[:2000]
        title, company = job.get("title", "the role"), job.get("company", "the company")
        emphasis = [e for e in (getattr(prof, "emphasis", []) or []) if e]
        system = (
            "You are an expert cover-letter writer. Return ONLY a JSON array of 3-4 paragraph "
            "strings (no greeting, no signature, no markdown). Tailor to the job using the "
            "candidate's real background; never invent experience."
            + (f" The OPENING paragraph must lead with the candidate's strength in "
               f"{', '.join(emphasis)} and their fit for the target role; frame "
               f"Kafka/Spark/Kubernetes as ML-infrastructure/MLOps work where relevant."
               if emphasis else "")
        )
        ask = (f"Write the body paragraphs of a cover letter.\nCANDIDATE: {prof.name}\n"
               f"SKILLS: {', '.join(prof.skills)}\n"
               + (f"LEAD WITH: {', '.join(emphasis)}\n" if emphasis else "")
               + f"BACKGROUND: {prof.resume_text}\n\n"
               f"JOB: {title} at {company}\nJOB DESCRIPTION:\n{jd}")
        msg = client.messages.create(model="claude-haiku-4-5", max_tokens=1400, system=system,
                                     messages=[{"role": "user", "content": ask}])
        raw = "".join(b.text for b in msg.content if b.type == "text").strip()
        raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
        raw = raw[raw.find("["): raw.rfind("]") + 1]
        arr = json.loads(raw)
        return [str(p) for p in arr if str(p).strip()] or None
    except Exception:
        return None


# ===========================================================================
# Word (.docx) renderers
# ===========================================================================
def _docx_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8); p.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RGBColor(*NAVY)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1F4E79")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def _docx_entry(doc, e):
    """One experience/project entry: 'Title | Company' left + 'location · dates' right, bullets."""
    p = doc.add_paragraph()
    p.space_after = Pt(1)
    # right tab stop at the page width for the dates
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(e["title"]); r.bold = True; r.font.size = Pt(10.5)
    if e["company"]:
        sep = p.add_run("  |  "); sep.bold = True; sep.font.size = Pt(10.5)
        c = p.add_run(e["company"]); c.bold = True; c.font.size = Pt(10.5); c.font.color.rgb = RGBColor(*NAVY)
    right = " · ".join(x for x in [e.get("location", ""), e.get("dates", "")] if x)
    if right:
        rr = p.add_run("\t" + right); rr.font.size = Pt(9.5); rr.font.color.rgb = RGBColor(*GRAY)
    for b in e["bullets"]:
        bp = doc.add_paragraph(b, style="List Bullet"); bp.space_after = Pt(1)
        for run in bp.runs:
            run.font.size = Pt(10)


def resume_docx_bytes(data: dict) -> bytes:
    doc = Document()
    for m in doc.sections:
        m.top_margin = Inches(0.5); m.bottom_margin = Inches(0.5)
        m.left_margin = Inches(0.7); m.right_margin = Inches(0.7)
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    name_p = doc.add_paragraph(); name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = name_p.add_run(data["name"].upper()); nr.bold = True
    nr.font.size = Pt(20); nr.font.color.rgb = RGBColor(*NAVY)
    if data["contact"]:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(data["contact"]); cr.font.size = Pt(9.5); cr.font.color.rgb = RGBColor(*GRAY)

    def _profile():
        if data["summary"]:
            _docx_heading(doc, "Profile"); doc.add_paragraph(data["summary"])

    def _publications():
        if data.get("publications"):
            _docx_heading(doc, "Publications")
            for pub in data["publications"]:
                bp = doc.add_paragraph(pub, style="List Bullet"); bp.space_after = Pt(1)
                for run in bp.runs:
                    run.font.size = Pt(10)

    def _exp(title, entries):
        if entries:
            _docx_heading(doc, title)
            for e in entries:
                _docx_entry(doc, e)

    def _skills():
        if data["skills"]:
            _docx_heading(doc, "Skills")
            for grp, items in data["skills"].items():
                p = doc.add_paragraph(); p.space_after = Pt(2)
                if grp and grp != "Skills":
                    g = p.add_run(f"{grp}: "); g.bold = True; g.font.size = Pt(10)
                t = p.add_run(" · ".join(items)); t.font.size = Pt(10)

    def _education():
        if data["education"]:
            _docx_heading(doc, "Education")
            for ed in data["education"]:
                p = doc.add_paragraph(); p.space_after = Pt(1)
                p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
                head = ed["degree"] + (f" · {ed['school']}" if ed["school"] else "")
                r = p.add_run(head); r.bold = True; r.font.size = Pt(10.5)
                if ed["dates"]:
                    rr = p.add_run("\t" + ed["dates"]); rr.font.size = Pt(9.5); rr.font.color.rgb = RGBColor(*GRAY)
                if ed["coursework"]:
                    cp = doc.add_paragraph(); cr = cp.add_run("Coursework: " + ed["coursework"])
                    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(*GRAY)

    render = {"profile": _profile, "publications": _publications,
              "experience": lambda: _exp("Professional Experience", data["experience"]),
              "projects": lambda: _exp("Projects", data["projects"]),
              "skills": _skills, "education": _education}
    for sec in data.get("section_order", ["profile", "experience", "projects", "skills", "education"]):
        render.get(sec, lambda: None)()

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def cover_letter_docx_bytes(data: dict) -> bytes:
    doc = Document()
    for m in doc.sections:
        m.top_margin = Inches(0.9); m.bottom_margin = Inches(0.9)
        m.left_margin = Inches(1.0); m.right_margin = Inches(1.0)
    st = doc.styles["Normal"]; st.font.name = "Georgia"; st.font.size = Pt(11)

    np_ = doc.add_paragraph(); nr = np_.add_run(data["name"].upper())
    nr.bold = True; nr.font.size = Pt(18); nr.font.color.rgb = RGBColor(*NAVY)
    if data["contact"]:
        cp = doc.add_paragraph(); cr = cp.add_run(data["contact"])
        cr.font.size = Pt(10); cr.font.color.rgb = RGBColor(*GRAY)
    doc.add_paragraph("")
    doc.add_paragraph(data["greeting"])
    doc.add_paragraph("")
    for para in data["paragraphs"]:
        p = doc.add_paragraph(para); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.space_after = Pt(8)
    doc.add_paragraph(data["closing"])
    doc.add_paragraph(data["name"])
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ===========================================================================
# PDF renderers (fpdf2 — pure Python, no system deps)
# ===========================================================================
def _new_pdf(serif: bool = False):
    from fpdf import FPDF
    pdf = FPDF(orientation="P", unit="pt", format="letter")
    pdf.set_auto_page_break(True, margin=40)
    pdf.set_margins(left=50, top=44, right=50)
    pdf.add_page()
    pdf.set_font("Times" if serif else "Helvetica", size=10)
    return pdf


def _pdf_section(pdf, title):
    pdf.ln(7)
    pdf.set_font("Helvetica", "B", 10.5); pdf.set_text_color(*NAVY)
    pdf.cell(0, 13, _latin1(title.upper())); pdf.ln(13)
    y = pdf.get_y()
    pdf.set_draw_color(*NAVY); pdf.set_line_width(0.8)
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(4)


def _pdf_entry(pdf, e):
    y = pdf.get_y()
    pdf.set_xy(pdf.l_margin, y)
    pdf.set_font("Helvetica", "B", 10.5); pdf.set_text_color(20, 20, 20)
    pdf.cell(pdf.get_string_width(_latin1(e["title"])) + 1, 14, _latin1(e["title"]))
    if e["company"]:
        pdf.set_text_color(*NAVY)
        pdf.cell(1 + pdf.get_string_width(_latin1("  |  " + e["company"])), 14, _latin1("  |  " + e["company"]))
    right = " - ".join(x for x in [e.get("location", ""), e.get("dates", "")] if x)
    if right:
        pdf.set_xy(pdf.l_margin, y)
        pdf.set_font("Helvetica", "", 9.3); pdf.set_text_color(*GRAY)
        pdf.cell(pdf.epw, 14, _latin1(right), align="R")
    pdf.set_xy(pdf.l_margin, y + 15)
    for b in e["bullets"]:
        yy = pdf.get_y()
        pdf.set_fill_color(*NAVY)
        pdf.ellipse(pdf.l_margin + 3, yy + 4.5, 2.4, 2.4, style="F")
        pdf.set_xy(pdf.l_margin + 13, yy)
        pdf.set_font("Helvetica", "", 10); pdf.set_text_color(35, 35, 35)
        pdf.multi_cell(pdf.epw - 13, 13, _latin1(b), align="L")
        pdf.set_x(pdf.l_margin)
    pdf.ln(3)


def resume_pdf_bytes(data: dict) -> bytes:
    pdf = _new_pdf()
    pdf.set_font("Helvetica", "B", 19); pdf.set_text_color(*NAVY)
    pdf.cell(0, 24, _latin1(data["name"].upper()), align="C"); pdf.ln(24)
    if data["contact"]:
        pdf.set_font("Helvetica", "", 9.5); pdf.set_text_color(*GRAY)
        pdf.cell(0, 13, _latin1(data["contact"]), align="C"); pdf.ln(13)

    def _profile():
        if data["summary"]:
            _pdf_section(pdf, "Profile")
            pdf.set_font("Helvetica", "", 10); pdf.set_text_color(35, 35, 35)
            pdf.multi_cell(pdf.epw, 13, _latin1(data["summary"]), align="L")

    def _publications():
        if data.get("publications"):
            _pdf_section(pdf, "Publications")
            pdf.set_font("Helvetica", "", 10); pdf.set_text_color(35, 35, 35)
            for pub in data["publications"]:
                yy = pdf.get_y()
                pdf.set_fill_color(*NAVY); pdf.ellipse(pdf.l_margin + 3, yy + 4.5, 2.4, 2.4, style="F")
                pdf.set_xy(pdf.l_margin + 13, yy)
                pdf.multi_cell(pdf.epw - 13, 13, _latin1(pub))
                pdf.set_x(pdf.l_margin)

    def _exp(title, entries):
        if entries:
            _pdf_section(pdf, title)
            for e in entries:
                _pdf_entry(pdf, e)

    def _skills():
        if data["skills"]:
            _pdf_section(pdf, "Skills")
            pdf.set_font("Helvetica", "", 10); pdf.set_text_color(40, 40, 40)
            for grp, items in data["skills"].items():
                label = f"**{grp}:**  " if grp and grp != "Skills" else ""
                pdf.multi_cell(pdf.epw, 13, _latin1(label + " · ".join(items)), markdown=True)
                pdf.ln(1)

    def _education():
        if data["education"]:
            _pdf_section(pdf, "Education")
            for ed in data["education"]:
                y = pdf.get_y()
                pdf.set_font("Helvetica", "B", 10.5); pdf.set_text_color(20, 20, 20)
                head = ed["degree"] + (f"  ·  {ed['school']}" if ed["school"] else "")
                pdf.set_xy(pdf.l_margin, y); pdf.cell(pdf.epw * 0.72, 14, _latin1(head))
                if ed["dates"]:
                    pdf.set_xy(pdf.l_margin, y); pdf.set_font("Helvetica", "", 9.3); pdf.set_text_color(*GRAY)
                    pdf.cell(pdf.epw, 14, _latin1(ed["dates"]), align="R")
                pdf.set_xy(pdf.l_margin, y + 14)
                if ed["coursework"]:
                    pdf.set_font("Helvetica", "I", 9); pdf.set_text_color(*GRAY)
                    pdf.multi_cell(pdf.epw, 12, _latin1("Coursework: " + ed["coursework"]))

    render = {"profile": _profile, "publications": _publications,
              "experience": lambda: _exp("Professional Experience", data["experience"]),
              "projects": lambda: _exp("Projects", data["projects"]),
              "skills": _skills, "education": _education}
    for sec in data.get("section_order", ["profile", "experience", "projects", "skills", "education"]):
        render.get(sec, lambda: None)()
    return bytes(pdf.output())


def cover_letter_pdf_bytes(data: dict) -> bytes:
    pdf = _new_pdf(serif=True)
    pdf.set_font("Times", "B", 20); pdf.set_text_color(*NAVY)
    pdf.cell(0, 26, _latin1(data["name"].upper())); pdf.ln(28)
    if data["contact"]:
        pdf.set_font("Times", "", 10.5); pdf.set_text_color(*GRAY)
        pdf.cell(0, 14, _latin1(data["contact"])); pdf.ln(22)
    pdf.set_font("Times", "", 11); pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(pdf.epw, 15, _latin1(data["greeting"])); pdf.ln(6)
    for para in data["paragraphs"]:
        pdf.multi_cell(pdf.epw, 15, _latin1(para), align="J"); pdf.ln(8)
    pdf.multi_cell(pdf.epw, 15, _latin1(data["closing"])); pdf.ln(2)
    pdf.multi_cell(pdf.epw, 15, _latin1(data["name"]))
    return bytes(pdf.output())


# ===========================================================================
# Plain-text preview (shown in the UI)
# ===========================================================================
def resume_plaintext(data: dict) -> str:
    out = [data["name"].upper()]
    if data["contact"]:
        out.append(data["contact"])

    def _exp(title, entries):
        nonlocal out
        if entries:
            out.append(""); out.append(title.upper())
            for e in entries:
                hdr = e["title"] + (f" | {e['company']}" if e["company"] else "")
                right = " · ".join(x for x in [e.get("location", ""), e.get("dates", "")] if x)
                out.append(f"{hdr}{('   ' + right) if right else ''}")
                out += [f"  • {b}" for b in e["bullets"]]

    def _do(sec):
        nonlocal out
        if sec == "profile" and data["summary"]:
            out += ["", "PROFILE", data["summary"]]
        elif sec == "publications" and data.get("publications"):
            out += ["", "PUBLICATIONS"] + [f"  • {p}" for p in data["publications"]]
        elif sec == "experience":
            _exp("Professional Experience", data["experience"])
        elif sec == "projects":
            _exp("Projects", data["projects"])
        elif sec == "skills" and data["skills"]:
            out += ["", "SKILLS"]
            for grp, items in data["skills"].items():
                out.append((f"{grp}: " if grp and grp != "Skills" else "") + " · ".join(items))
        elif sec == "education" and data["education"]:
            out += ["", "EDUCATION"]
            for ed in data["education"]:
                out.append(ed["degree"] + (f" · {ed['school']}" if ed["school"] else "")
                           + (f"   {ed['dates']}" if ed["dates"] else ""))
                if ed["coursework"]:
                    out.append(f"  Coursework: {ed['coursework']}")

    for sec in data.get("section_order", ["profile", "experience", "projects", "skills", "education"]):
        _do(sec)
    return "\n".join(out)


def cover_plaintext(data: dict) -> str:
    parts = [data["name"].upper()]
    if data["contact"]:
        parts.append(data["contact"])
    parts += ["", data["greeting"], ""]
    parts += data["paragraphs"]
    parts += ["", data["closing"], data["name"]]
    return "\n".join(parts)


# ===========================================================================
# Public API — returns dict with both formats
# ===========================================================================
def generate_resume(job: pd.Series, prof, use_llm: bool = True) -> dict:
    data = build_resume_data(job, prof, use_llm)
    return {"docx": resume_docx_bytes(data), "pdf": resume_pdf_bytes(data),
            "text": resume_plaintext(data), "ats": data["ats"], "backend": data["backend"]}


def generate_cover_letter(job: pd.Series, prof, use_llm: bool = True) -> dict:
    data = build_cover_data(job, prof, use_llm)
    return {"docx": cover_letter_docx_bytes(data), "pdf": cover_letter_pdf_bytes(data),
            "text": cover_plaintext(data), "backend": data["backend"]}
